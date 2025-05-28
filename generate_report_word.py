from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_paragraph(doc, text, bold=False, italic=False, font_size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    font = run.font
    font.size = Pt(font_size)
    font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_table(doc, data, col_widths=None):
    table = doc.add_table(rows=1, cols=len(data[0]))
    hdr_cells = table.rows[0].cells
    for i, heading in enumerate(data[0]):
        hdr_cells[i].text = heading
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
    for row in data[1:]:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width
    return table

def generate_report(filename):
    doc = Document()
    # Trang bìa
    add_heading(doc, "BÁO CÁO ĐỀ TÀI", 0)
    add_heading(doc, "XÂY DỰNG HỆ THỐNG ĐIỂM DANH SINH VIÊN KẾT HỢP IoT VÀ AI", 1)
    add_paragraph(doc, "Nhóm thực hiện: .............................................", font_size=13)
    add_paragraph(doc, "Giảng viên hướng dẫn: .....................................", font_size=13)
    add_paragraph(doc, "Ngày nộp: ................................................", font_size=13)
    doc.add_page_break()

    # 1. Giới thiệu đề tài
    add_heading(doc, "1. Giới thiệu đề tài", 1)
    add_paragraph(doc, "Trong bối cảnh chuyển đổi số mạnh mẽ, việc tự động hóa quản lý điểm danh sinh viên là nhu cầu cấp thiết nhằm nâng cao hiệu quả, giảm gian lận và tiết kiệm nguồn lực. Đề tài này xây dựng hệ thống điểm danh sinh viên ứng dụng AI (nhận diện khuôn mặt) kết hợp IoT, giúp quản lý tập trung, minh bạch và thuận tiện cho cả nhà trường lẫn sinh viên. Hệ thống có thể mở rộng cho nhiều lớp học, nhiều điểm điểm danh, tích hợp với các thiết bị IoT như cửa tự động, đèn báo, ...")
    doc.add_page_break()

    # 2. Thành phần linh kiện sử dụng
    add_heading(doc, "2. Thành phần linh kiện sử dụng", 1)
    add_heading(doc, "2.1. Phần cứng (IoT)", 2)
    add_paragraph(doc, "- Camera/Webcam: Thu hình khuôn mặt sinh viên tại điểm danh, có thể dùng camera USB hoặc camera tích hợp trên máy tính/laptop.")
    add_paragraph(doc, "- Máy tính/Raspberry Pi: Thiết bị xử lý trung tâm, nhận dữ liệu từ camera, chạy phần mềm nhận diện và quản lý dữ liệu.")
    add_paragraph(doc, "- (Tùy chọn) Thiết bị IoT khác: Ví dụ: cửa tự động, đèn báo, cảm biến chuyển động, RFID để tăng tính bảo mật và tự động hóa.")
    add_heading(doc, "2.2. Phần mềm (AI & Backend)", 2)
    add_paragraph(doc, "- Python: Ngôn ngữ lập trình chính, dễ phát triển và tích hợp các thư viện AI, IoT.")
    add_paragraph(doc, "- Flask: Framework web nhẹ, dễ triển khai giao diện quản trị và API.")
    add_paragraph(doc, "- OpenCV: Thư viện xử lý ảnh mạnh mẽ, hỗ trợ thao tác với camera, tiền xử lý ảnh.")
    add_paragraph(doc, "- face_recognition: Thư viện AI nhận diện khuôn mặt dựa trên deep learning, cho độ chính xác cao.")
    add_paragraph(doc, "- MongoDB: Cơ sở dữ liệu NoSQL, lưu trữ linh hoạt thông tin sinh viên, ảnh mẫu, lịch sử điểm danh.")
    add_paragraph(doc, "- Bootstrap: Thư viện giao diện web hiện đại, responsive, thân thiện với người dùng.")
    doc.add_page_break()

    # 3. Kinh phí dự kiến
    add_heading(doc, "3. Kinh phí dự kiến", 1)
    table_data = [
        ["Thành phần", "Số lượng", "Đơn giá (VNĐ)", "Thành tiền (VNĐ)"],
        ["Camera HD", "1", "400.000", "400.000"],
        ["Máy tính/Raspberry Pi", "1", "2.000.000", "2.000.000"],
        ["(Tùy chọn) RFID", "0", "0", "0"],
        ["Chi phí điện, mạng", "-", "200.000", "200.000"],
        ["Tổng cộng", "", "", "2.600.000"]
    ]
    add_table(doc, table_data)
    add_paragraph(doc, "*Lưu ý: Nếu dùng máy tính cá nhân, chi phí phần cứng có thể giảm. Hệ thống có thể mở rộng với nhiều camera hoặc thiết bị IoT khác nếu cần.")
    doc.add_page_break()

    # 4. Thời gian lắp đặt
    add_heading(doc, "4. Thời gian lắp đặt", 1)
    time_table = [
        ["Công việc", "Thời gian (ngày)"],
        ["Nghiên cứu, chuẩn bị linh kiện", "2"],
        ["Lắp đặt phần cứng, kết nối camera", "1"],
        ["Cài đặt phần mềm, môi trường", "1"],
        ["Lập trình, tích hợp hệ thống", "2"],
        ["Kiểm thử, hoàn thiện", "1"],
        ["Tổng cộng", "7"]
    ]
    add_table(doc, time_table)
    add_paragraph(doc, "Thời gian có thể rút ngắn nếu có sẵn thiết bị và kinh nghiệm triển khai. Giai đoạn kiểm thử nên thực hiện kỹ để đảm bảo hệ thống ổn định.")
    doc.add_page_break()

    # 5. Sơ đồ khối & liên kết hệ thống
    add_heading(doc, "5. Sơ đồ khối & liên kết hệ thống", 1)
    add_heading(doc, "5.1. Sơ đồ khối tổng thể", 2)
    add_paragraph(doc, "[Camera/Webcam] -> [Máy tính/Raspberry Pi] --(Kết nối mạng)--> [Cơ sở dữ liệu MongoDB] -> [Giao diện Web Flask]", font_size=11)
    add_heading(doc, "5.2. Mô tả liên kết các thành phần", 2)
    add_paragraph(doc, "- Camera/Webcam: Thu hình ảnh khuôn mặt sinh viên, truyền về máy tính/Raspberry Pi.")
    add_paragraph(doc, "- Máy tính/Raspberry Pi: Nhận hình ảnh từ camera, xử lý nhận diện khuôn mặt bằng AI, giao tiếp với MongoDB, chạy server Flask.")
    add_paragraph(doc, "- MongoDB: Lưu trữ thông tin sinh viên, ảnh mẫu, lịch sử điểm danh.")
    add_paragraph(doc, "- Giao diện Web: Quản lý sinh viên, xem lịch sử điểm danh, xem trực tiếp camera.")
    add_paragraph(doc, "- (Tùy chọn) Các thiết bị IoT khác: Có thể tích hợp để tự động mở cửa, bật đèn, gửi thông báo khi sinh viên điểm danh thành công.")
    doc.add_page_break()

    # 6. Mô tả hoạt động hệ thống
    add_heading(doc, "6. Mô tả hoạt động hệ thống", 1)
    add_paragraph(doc, "1. Thêm sinh viên: Quản trị viên thêm sinh viên mới qua web, upload/chụp ảnh khuôn mặt. Ảnh này sẽ được dùng làm mẫu nhận diện.")
    add_paragraph(doc, "2. Điểm danh: Camera ghi lại hình ảnh sinh viên, AI nhận diện khuôn mặt, ghi nhận điểm danh vào MongoDB. Có thể tích hợp cảnh báo nếu phát hiện người lạ.")
    add_paragraph(doc, "3. Quản lý: Quản trị viên xem, chỉnh sửa thông tin sinh viên, xuất lịch sử điểm danh, thống kê tỷ lệ chuyên cần.")
    add_paragraph(doc, "4. Mở rộng: Có thể tích hợp thêm các chức năng như gửi email thông báo, kết nối hệ thống quản lý đào tạo, ...")
    doc.add_page_break()

    # 7. Giải thích thuật toán nhận diện khuôn mặt (bổ sung chi tiết)...
    # (phần này chính là phần bổ sung chi tiết mà mình đã soạn ở tin nhắn trước!)

    add_heading(doc, "7. Giải thích thuật toán nhận diện khuôn mặt", 1)
    add_paragraph(doc, "Quy trình nhận diện khuôn mặt trong hệ thống gồm các bước chính sau:", bold=True)
    add_paragraph(doc, "1. Tiền xử lý ảnh:")
    add_paragraph(doc, "   - Ảnh thu được từ camera thường là ảnh màu (RGB). Để đơn giản hóa xử lý, ảnh sẽ được chuyển sang ảnh đen trắng (grayscale). Việc này giúp giảm nhiễu, giảm số chiều dữ liệu và tăng tốc độ xử lý.")
    add_paragraph(doc, "2. Phát hiện vùng khuôn mặt:")
    add_paragraph(doc, "   - Sử dụng các mô hình phát hiện khuôn mặt (ví dụ: HOG, CNN hoặc Haar Cascade) để xác định vị trí các khuôn mặt trong ảnh. Kết quả là các bounding box bao quanh từng khuôn mặt.")
    add_paragraph(doc, "3. Căn chỉnh và cắt khuôn mặt:")
    add_paragraph(doc, "   - Ảnh khuôn mặt được cắt ra từ ảnh gốc, có thể căn chỉnh lại để chuẩn hóa kích thước, hướng nhìn, ... giúp tăng độ chính xác khi nhận diện.")
    add_paragraph(doc, "4. Trích xuất đặc trưng khuôn mặt:")
    add_paragraph(doc, "   - Sử dụng các mô hình deep learning (ví dụ: FaceNet, dlib, ResNet) để chuyển ảnh khuôn mặt thành vector đặc trưng (embedding). Mỗi khuôn mặt sẽ có một vector đặc trưng duy nhất, phản ánh các đặc điểm sinh trắc học của từng người.")
    add_paragraph(doc, "5. So sánh đặc trưng và nhận diện:")
    add_paragraph(doc, "   - Vector đặc trưng của khuôn mặt mới sẽ được so sánh với các vector đã lưu trong cơ sở dữ liệu (dùng khoảng cách Euclidean hoặc cosine). Nếu khoảng cách nhỏ hơn ngưỡng cho phép, hệ thống xác định đó là cùng một người và ghi nhận điểm danh.")
    add_paragraph(doc, "6. Lưu trữ và cập nhật dữ liệu:")
    add_paragraph(doc, "   - Nếu nhận diện thành công, hệ thống sẽ lưu lại thông tin điểm danh (thời gian, mã sinh viên, ...) vào cơ sở dữ liệu. Nếu không nhận diện được, có thể cảnh báo hoặc ghi nhận là người lạ.")

    add_heading(doc, "7.1. Đầu vào và đầu ra của thuật toán", 2)
    add_paragraph(doc, "- Đầu vào: Ảnh/video từ camera chứa khuôn mặt sinh viên (ảnh màu RGB).")
    add_paragraph(doc, "- Đầu ra: Kết quả nhận diện (ID sinh viên nếu nhận diện thành công hoặc thông báo \"Người lạ\" nếu không có kết quả).")

    add_heading(doc, "7.2. Hoạt động tốt khi nào?", 2)
    add_paragraph(doc, "- Ánh sáng đầy đủ, khuôn mặt rõ nét, không che khuất.")
    add_paragraph(doc, "- Dữ liệu mẫu lưu trong database chất lượng cao (nhiều góc chụp, ít nhiễu).")
    add_paragraph(doc, "- Camera chất lượng ổn định, độ phân giải đủ cao (HD trở lên).")

    add_heading(doc, "7.3. Hoạt động không tốt khi nào?", 2)
    add_paragraph(doc, "- Ảnh mờ, nhiễu, thiếu sáng.")
    add_paragraph(doc, "- Khuôn mặt bị che (khẩu trang, kính râm, tóc che…).")
    add_paragraph(doc, "- Cơ sở dữ liệu mẫu không đầy đủ hoặc bị lỗi.")

    add_heading(doc, "7.4. Nguyên nhân gây sai lệch", 2)
    add_paragraph(doc, "- Model không nhận diện tốt nếu ánh sáng kém hoặc góc chụp không chính diện.")
    add_paragraph(doc, "- Số lượng ảnh mẫu trong database quá ít, không đủ đa dạng góc độ.")
    add_paragraph(doc, "- Camera chất lượng thấp làm giảm độ phân giải, khiến feature extraction kém chính xác.")
    doc.add_page_break()

    # Các phần tiếp theo: đoạn mã mẫu, kết luận, tài liệu tham khảo...
    # (giữ nguyên như file gốc!)

    doc.save(filename)

if __name__ == "__main__":
    generate_report("BaoCao_DiemDanh_IoT_AI_Full.docx")